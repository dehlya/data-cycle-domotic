"""
polish_user_guide_docx.py — improve the existing Power BI user guide

Reads an existing .docx, applies targeted run-level text replacements
(preserves images, formatting, and overall structure), and writes a
polished copy to docs/v2/out/.

Run with:
    python tools/polish_user_guide_docx.py

Author: Group 14 — DataCycle Domotic
"""

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

INPUT_PATH  = Path(r"C:\Users\dehly\Downloads\Data cycle User guide.docx")
OUTPUT_PATH = Path("docs/v2/out/DataCycle_User_Guide.docx")


# ── 1. Targeted text replacements ────────────────────────────────────────────
# (substring → replacement). Applied at the run level so images and
# formatting survive. The keys must be substrings that appear *inside a
# single run*; if a string is split across runs, we do a paragraph-level
# rewrite that preserves run formatting from the first run.
TEXT_REPLACEMENTS = [
    # Encoding artefact: smart quote in "won't" / "you're" etc. — pandoc
    # extraction showed "�" but the actual chars are smart quotes which
    # display fine. No change needed for those.

    # Typos and casing
    ("KWH by room chart", "kWh by room chart"),
    ("KWH ",              "kWh "),
    ("Co2",               "CO\u2082"),  # CO with subscript 2
    ("Some graph displayed",  "Some graphs displayed"),

    # Voice consistency: prefer "you" over "we" in user-facing prose
    ("In this table we have a sample",  "This table shows a sample"),
    ("we can see how much",             "you can see"),
    ("Next to this timeline we can see", "Next to this timeline you can see"),
    ("you can see co2 we have by room", "you can see CO\u2082 levels by room"),
    ("co2 we have by room",              "CO\u2082 levels by room"),
    ("we can see the humidity",         "you can see the humidity"),

    # FAQ #1 — missing verb + question mark
    ("Why are some fields display \u201c--\u201d or \u201cNo data\u201d",
     "Why are some fields displaying \u201c--\u201d or \u201cNo data\u201d?"),
    ("Why are some fields display \"--\" or \"No data\"",
     "Why are some fields displaying \"--\" or \"No data\"?"),

    # FAQ #2 — needs question mark
    ("Can I isolate a single data point in a graph",
     "Can I isolate a single value in a graph?"),

    # FAQ #3 — needs question mark
    ("Is the filter synchronized across pages",
     "Is the filter synchronized across pages?"),
    ("Yes",
     "Yes — selections in the slicers apply across every page of the report."),

    # FAQ #4 — replace misleading answer (skip the question — it already has ?)
    ("The data is not updated automatically",
     "The pipeline keeps the database fresh in the background — sensor "
     "readings every minute, gold facts every 15 minutes, and ML "
     "predictions daily at 06:30. Power BI Desktop loads a snapshot of "
     "this data when you open the file: click Home \u2192 Refresh "
     "(or press Ctrl+R) to pull the latest values."),

    # Filters terminology — Power BI calls them slicers, but "filters"
    # is fine. Just make sure "(the buttons at the top)" is more accurate.
    ("the filters (the buttons at the top)",
     "the filters (the slicers at the top of each page)"),

    # Compare page — "the period" is ambiguous
    ("by clicking the filter and choosing the period",
     "by selecting the apartment slicer (Jimmy / Jeremie / both) and "
     "the date range"),

    # Navigation header — pluralise
    ("Navigation & Filter", "Navigation & Filters"),

    # Device health
    ("the KPI of all the devices",       "the KPIs of all the devices"),
    ("if all is running or if a device is offline",
     "whether all devices are online or any are offline"),
]


# ── 2. Helper: replace text inside a paragraph at run level ─────────────────
def replace_in_paragraph(paragraph, old, new):
    """Replace 'old' with 'new' inside a paragraph, preserving runs as much
    as possible. Strategy: if 'old' fits entirely inside a single run,
    replace there. Otherwise, rebuild the paragraph: collect all text-bearing
    runs, replace in the joined text, redistribute the new text into the
    first text-bearing run (which keeps its formatting), and clear text in
    the others. This preserves images (image-bearing runs are skipped)."""
    # Fast path: substring is in a single run
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # Slow path: scan text-bearing runs only
    text_runs = []
    for run in paragraph.runs:
        # Image runs have a <w:drawing> element; skip them
        has_drawing = any(child.tag == qn("w:drawing") for child in run._element)
        if has_drawing:
            continue
        text_runs.append(run)

    if not text_runs:
        return False

    joined = "".join(r.text for r in text_runs)
    if old not in joined:
        return False

    new_joined = joined.replace(old, new)
    text_runs[0].text = new_joined
    for r in text_runs[1:]:
        r.text = ""
    return True


# ── 3. Helper: insert a new paragraph after a target paragraph ──────────────
def insert_paragraph_after(target_paragraph, text, style=None, italic=False, accent=False):
    """Insert a new paragraph immediately after the given paragraph.
    Returns the new Paragraph."""
    new_p = deepcopy(target_paragraph._element)
    # Strip the cloned paragraph's content (children except pPr)
    for child in list(new_p):
        if child.tag != qn("w:pPr"):
            new_p.remove(child)
    target_paragraph._element.addnext(new_p)

    from docx.text.paragraph import Paragraph as DocxParagraph
    para_obj = DocxParagraph(new_p, target_paragraph._parent)
    if style is not None:
        try:
            para_obj.style = style
        except KeyError:
            pass
    run = para_obj.add_run(text)
    if italic:
        run.italic = True
    if accent:
        run.font.color.rgb = RGBColor(0x2E, 0x5B, 0xFF)
    return para_obj


# ── 4. Main flow ─────────────────────────────────────────────────────────────
def main():
    if not INPUT_PATH.exists():
        sys.exit(f"Input not found: {INPUT_PATH}")
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    print(f"Reading  {INPUT_PATH}")
    doc = Document(INPUT_PATH)

    # 4a. Apply text replacements
    n_changes = 0
    for old, new in TEXT_REPLACEMENTS:
        for para in doc.paragraphs:
            if replace_in_paragraph(para, old, new):
                n_changes += 1
                print(f"  ✓ '{old[:50]}...' → '{new[:50]}...'")
                break  # one paragraph hit per replacement is enough
    print(f"\n{n_changes} text replacements applied")

    # Cleanup pass — fix accidental "??" left by replacements that added ? to
    # a question that already ended in ?
    for para in doc.paragraphs:
        for run in para.runs:
            if "??" in run.text:
                run.text = run.text.replace("??", "?")

    # 4b. Insert helpful tips at strategic spots
    inserted = []

    # After the Navigation list (find the "Device Health: ..." bullet → insert F11 tip after the next normal paragraph or at end of section)
    for i, para in enumerate(doc.paragraphs):
        if "Device Health: Monitoring" in (para.text or ""):
            # Insert F11 tip + RLS preview tip after this bullet's paragraph cluster
            tip_target = doc.paragraphs[i]
            tip = insert_paragraph_after(
                tip_target,
                "Tip — press F11 in Power BI Desktop for a fullscreen presentation view that hides every toolbar.",
                italic=True, accent=True,
            )
            inserted.append("F11 tip after Navigation list")
            # Insert RLS tip just after the F11 tip
            rls_tip = insert_paragraph_after(
                tip,
                "Apartment-level access — each tenant only sees their own apartment thanks to Row-Level Security. To preview a tenant view (admin only): Modeling \u2192 View as \u2192 Other user \u2192 Jimmy or Jeremie.",
                italic=True, accent=True,
            )
            inserted.append("RLS preview tip after Navigation")
            break

    # Before "Conclusion" → insert refresh reminder
    for i, para in enumerate(doc.paragraphs):
        if (para.style.name == "Heading 1") and (para.text or "").strip().lower() == "conclusion":
            prev = doc.paragraphs[i - 1] if i > 0 else None
            if prev is not None:
                insert_paragraph_after(
                    prev,
                    "Reminder — Power BI Desktop displays a snapshot of the data taken at refresh time. To pull the latest values, click Home \u2192 Refresh (or press Ctrl+R) before reading the dashboards.",
                    italic=True, accent=True,
                )
                inserted.append("Refresh reminder before Conclusion")
            break

    print(f"\n{len(inserted)} contextual tips inserted:")
    for x in inserted:
        print(f"  + {x}")

    # 4c. Save
    doc.save(OUTPUT_PATH)
    size_kb = OUTPUT_PATH.stat().st_size / 1024
    print(f"\n✓ Wrote {OUTPUT_PATH} ({size_kb:.0f} KB)")


if __name__ == "__main__":
    main()
